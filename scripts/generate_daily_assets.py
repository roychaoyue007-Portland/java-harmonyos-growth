import argparse
import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


DEFAULT_PROJECT = Path(r"E:\Codex\projects\java-harmonyos-growth")


def read_config(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\n")


def append_section_once(path: Path, marker: str, section: str) -> None:
    existing = path.read_text(encoding="utf-8") if path.exists() else ""
    if marker not in existing:
        write_text(path, existing.rstrip() + "\n\n" + section.strip() + "\n")


def list_lines(items):
    return "\n".join(f"- {item}" for item in items)


def numbered_lines(items):
    return "\n".join(f"{idx}. {item}" for idx, item in enumerate(items, 1))


def code_block(lines, lang="text"):
    return "```" + lang + "\n" + "\n".join(lines) + "\n```"


def build_notes_text(cfg: dict) -> str:
    day = cfg["day"]
    title = cfg["title"]
    concepts = cfg["concepts"]
    commands = cfg.get("commands", [])

    original_answers = cfg.get("original_answers", [])
    corrected_answers = cfg.get("corrected_answers", [])

    original_block = ""
    if original_answers:
        original_block = "\n## 我的原始验收答案\n\n" + numbered_lines(original_answers) + "\n"

    corrected_block = ""
    if corrected_answers:
        corrected_block = "\n## 我的验收答案整理\n\n"
        for idx, item in enumerate(corrected_answers, 1):
            corrected_block += f"{idx}. {item['answer']}\n"

    command_block = ""
    if commands:
        command_block = "\n\n" + code_block(commands, cfg.get("command_language", "powershell"))

    return f"""# {day} - {title}

日期：{cfg["date"]}

建议学习时长：{cfg.get("duration", "3 到 4 小时")}。

今天不学习 {cfg.get("avoid_topics", "后面阶段内容")}。

今天只解决一个问题：

{code_block([cfg["core_question"]])}

## 阶段 1：概念理解，60 分钟

理解 {len(concepts)} 个基础词：

{list_lines(concepts)}

每个词都要能回答：

- 它是什么？
- 它解决什么问题？
- 我在哪里会遇到？
- 它和其他几个词有什么关系？
- 我有没有亲手观察过？

## 阶段 2：核心关系，30 分钟

今天要先记住这条关系：

{code_block(cfg["core_relations"])}

{cfg["core_explanation"]}

## 阶段 3：动手观察，60 分钟

{cfg["hands_on_intro"]}{command_block}

观察重点：

{list_lines(cfg["observe_points"])}

如果出现报错，可以从这些方向排查：

{numbered_lines(cfg["troubleshooting"])}

## 阶段 4：今日笔记，40 分钟

回答下面问题：

{numbered_lines(cfg["note_questions"])}

## 阶段 5：技术词典，30 分钟

更新：

- `dictionary/tech-dictionary.md`

记录 {len(concepts)} 个词：

{list_lines(concepts)}

## 阶段 6：验收测试，30 到 40 分钟

回答：

{numbered_lines(cfg["acceptance_questions"])}
{original_block}{corrected_block}
## 补充理解

{cfg.get("extra_explanation", "")}

## 阶段 7：复盘，10 到 20 分钟

写下：

{list_lines(cfg["recap"])}

## 今日通过标准

学完后我应该能做到：

{list_lines(cfg["pass_criteria"])}
"""


def build_answer_text(cfg: dict) -> str:
    lines = [f"# {cfg['day']}：{cfg['title']}", ""]
    originals = cfg.get("original_answers", [])
    if originals:
        lines += ["## 我的原始答案", ""]
        lines.extend(f"{idx}. {answer}" for idx, answer in enumerate(originals, 1))
        lines.append("")

    corrected = cfg.get("corrected_answers", [])
    if corrected:
        lines += ["## 修正版答案", ""]
        for idx, item in enumerate(corrected, 1):
            lines += [f"### {idx}. {item['question']}", "", item["answer"], ""]

    memory = cfg.get("memory_lines", [])
    if memory:
        lines += ["## 最终记忆版", "", code_block(memory), ""]

    if cfg.get("conclusion"):
        lines += ["## 今日结论", "", cfg["conclusion"], ""]

    return "\n".join(lines)


def build_dictionary_section(cfg: dict) -> str:
    lines = [f"## {cfg['day']}"]
    for entry in cfg["dictionary_entries"]:
        lines += [
            "",
            f"### 词：{entry['term']}",
            "",
            f"一句话解释：{entry['one_line']}",
            "",
            f"它解决什么问题：{entry['solves']}",
            "",
            f"我在哪里会遇到：{entry['where']}",
            "",
            f"相关词：{entry['related']}",
            "",
            f"我是否亲手做过：{entry['done']}",
        ]
    return "\n".join(lines) + "\n"


def build_english_section(cfg: dict) -> str:
    lines = [f"## {cfg['day']}", "", "| English | 中文 | 例子 |", "| --- | --- | --- |"]
    for row in cfg["english_rows"]:
        lines.append(f"| {row['english']} | {row['chinese']} | {row['example']} |")
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11.5, "333333"),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build_docx(cfg: dict, path: Path) -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{cfg['day']}：{cfg['title']}")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"日期：{cfg['date']}    主题：{cfg['title']}").italic = True

    originals = cfg.get("original_answers", [])
    if originals:
        doc.add_heading("一、我的原始答案（保留）", level=1)
        for answer in originals:
            doc.add_paragraph(answer, style="List Number")

    corrected = cfg.get("corrected_answers", [])
    if corrected:
        doc.add_heading("二、修正版答案", level=1)
        for idx, item in enumerate(corrected, 1):
            doc.add_heading(f"{idx}. {item['question']}", level=2)
            doc.add_paragraph(item["answer"])

    memory = cfg.get("memory_lines", [])
    if memory:
        doc.add_heading("三、今天的核心关系", level=1)
        for line in memory:
            add_bullet(doc, line)

    doc.add_heading("四、技术词典", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, label in enumerate(["词", "一句话解释", "相关词"]):
        table.rows[0].cells[idx].text = label
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for entry in cfg["dictionary_entries"]:
        row = table.add_row().cells
        row[0].text = entry["term"]
        row[1].text = entry["one_line"]
        row[2].text = entry["related"]

    doc.add_heading("五、英文词汇", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, label in enumerate(["English", "中文", "例子"]):
        table.rows[0].cells[idx].text = label
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    for entry in cfg["english_rows"]:
        row = table.add_row().cells
        row[0].text = entry["english"]
        row[1].text = entry["chinese"]
        row[2].text = entry["example"]

    if cfg.get("conclusion"):
        doc.add_heading("六、今日结论", level=1)
        doc.add_paragraph(cfg["conclusion"])

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate daily study notes, answers, dictionary, English vocabulary, and optional DOCX.")
    parser.add_argument("--config", required=True, help="UTF-8 JSON config file.")
    parser.add_argument("--project", default=str(DEFAULT_PROJECT), help="Project root.")
    parser.add_argument("--outputs", default="", help="Optional external output folder for a copy of the answer Markdown.")
    parser.add_argument("--docx", action="store_true", help="Also generate DOCX files. Markdown is the default output.")
    args = parser.parse_args()

    cfg = read_config(Path(args.config))
    project = Path(args.project)
    day_dir = cfg["day"].lower().replace(" ", "-")
    note_slug = cfg["note_slug"]

    notes = build_notes_text(cfg)
    answers = build_answer_text(cfg)
    dict_section = build_dictionary_section(cfg)
    english_section = build_english_section(cfg)

    notes_path = project / "notes" / f"{note_slug}.md"
    practice_note_path = project / "practice" / day_dir / "笔记.md"
    answer_md_path = project / "practice" / day_dir / "answer" / "验收测试答案-修正版.md"
    answer_docx_path = project / "practice" / day_dir / "answer" / "验收测试答案.docx"

    write_text(notes_path, notes)
    write_text(practice_note_path, notes)
    write_text(answer_md_path, answers)
    append_section_once(project / "dictionary" / "tech-dictionary.md", f"## {cfg['day']}", dict_section)
    append_section_once(project / "english" / "vocabulary.md", f"## {cfg['day']}", english_section)

    print(f"WROTE {notes_path}")
    print(f"WROTE {practice_note_path}")
    print(f"WROTE {answer_md_path}")

    if args.docx:
        build_docx(cfg, answer_docx_path)
        print(f"WROTE {answer_docx_path}")

    if args.outputs:
        output_md = Path(args.outputs) / f"{cfg['day']} 验收测试答案.md"
        write_text(output_md, answers)
        print(f"WROTE {output_md}")
        if args.docx:
            output_docx = Path(args.outputs) / f"{cfg['day']} 验收测试答案.docx"
            build_docx(cfg, output_docx)
            print(f"WROTE {output_docx}")


if __name__ == "__main__":
    main()
